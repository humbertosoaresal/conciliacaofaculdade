# utils.py
import pandas as pd
from datetime import datetime
from io import BytesIO
import re

def safe_parse_date(date_str, default_date):
    """Tenta converter uma string para data, retornando uma data padrão em caso de falha."""
    if not date_str:
        return default_date
    try:
        # Adicionado dayfirst=True para garantir que o formato DD/MM/YYYY seja priorizado
        return pd.to_datetime(date_str, dayfirst=True).date()
    except (ValueError, TypeError):
        return default_date

def to_excel(df: pd.DataFrame) -> bytes:
    """Converte um DataFrame para um arquivo Excel em memória."""
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Relatorio')
    return output.getvalue()

def formatar_dataframe_para_exibicao(df, colunas_moeda):
    """Formata colunas de moeda e data para exibição no Streamlit."""
    df_display = df.copy()
    for col in colunas_moeda:
        if col in df_display.columns:
            df_display[col] = df_display[col].apply(lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
    
    if 'Data Lançamento' in df_display.columns:
        df_display['Data Lançamento'] = pd.to_datetime(df_display['Data Lançamento']).dt.strftime('%d/%m/%Y')
    
    return df_display

def convert_df_to_csv(df):
    """Converte um DataFrame para CSV em memória."""
    return df.to_csv(index=False, sep=';', decimal=',').encode('utf-8')

def create_word_report(data):
    """Cria um relatório simples no formato Word (placeholder)."""
    # Esta é uma função placeholder. A implementação real requer a biblioteca python-docx.
    try:
        from docx import Document
        document = Document()
        document.add_heading('Relatório de Conciliação', 0)
        
        for key, value in data.items():
            document.add_heading(key, level=1)
            if isinstance(value, pd.DataFrame):
                # Adiciona tabela
                t = document.add_table(rows=1, cols=len(value.columns))
                for i, col_name in enumerate(value.columns):
                    t.cell(0, i).text = col_name
                for index, row in value.iterrows():
                    row_cells = t.add_row().cells
                    for i, cell_value in enumerate(row):
                        row_cells[i].text = str(cell_value)
            else:
                document.add_paragraph(str(value))
        
        bio = BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()
    except ImportError:
        # Fallback se a biblioteca não estiver instalada
        return "Relatório (instale python-docx para formato Word)".encode('utf-8')
    except Exception as e:
        return f"Erro ao gerar relatório: {e}".encode('utf-8')

def extrair_conta_ofx_bruta(file_bytes: bytes) -> str:
    """Extrai o valor bruto da tag <ACCTID> de um arquivo OFX."""
    encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings_to_try:
        try:
            content = file_bytes.decode(encoding)
            match = re.search(r'<ACCTID>([^<]+)', content)
            if match:
                return match.group(1).strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception:
            pass

    return ''

def normalizar_numero(valor, is_conta_longa=False, is_conta_cadastro=False):
    """
    Normaliza um número de conta ou agência, removendo caracteres não numéricos.
    - is_conta_longa: Para contas que incluem agência (formato Sicredi).
    - is_conta_cadastro: Para contas do arquivo de cadastro que podem ter 'X' ou '-'.
    """
    if not isinstance(valor, str):
        valor = str(valor)
    
    if is_conta_cadastro:
        # Remove o dígito verificador se presente (ex: 1234-5 -> 1234)
        valor = valor.split('-')[0].split('X')[0]

    # Extrai todos os dígitos
    numeros = re.sub(r'\D', '', valor)
    
    if is_conta_longa:
        # Para contas no formato Sicredi (agência + conta), retorna os dígitos
        return numeros
    else:
        # Para outros casos, remove zeros à esquerda
        return numeros.lstrip('0') if numeros else '0'

def normalizar_chave_ofx(chave: str) -> str:
    """
    Centraliza a regra de normalização da chave OFX.
    Regra: 4 primeiros dígitos (agência) + restante da conta sem zeros à esquerda.
    Exemplo: '22050000000642886' -> '2205642886'
    """
    if not isinstance(chave, str) or len(chave) < 5:
        return chave # Retorna a chave original se for inválida
    
    chave_limpa = re.sub(r'\D', '', chave) # Garante que só temos dígitos
    agencia = chave_limpa[:4]
    conta = chave_limpa[4:].lstrip('0')
    
    return agencia + conta
